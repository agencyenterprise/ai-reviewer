"""Unit tests for `lib.services.docx.drawing_rasterization` — no LibreOffice.

The rendering leg (LibreOffice PDF export + pypdfium2) is covered by the
integration tests; everything here exercises the pure structure work: which
drawings are selected, how the render-source copy is built, and how rendered
images are spliced back into the document.
"""

import asyncio
import io
import os
import tempfile
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from PIL import Image

from lib.services.docx.drawing_rasterization import (
    CHART_GRAPHIC_URI,
    PICTURE_GRAPHIC_URI,
    _render_pdf_pages,
    _renderable_paragraphs,
    _replace_drawing_with_image,
    _write_render_source_docx,
    rasterize_docx_drawings,
)

MODULE = "lib.services.docx.drawing_rasterization"


# --- fixture helpers -------------------------------------------------------


def _chart_graphic_data() -> str:
    return (
        f'<a:graphicData uri="{CHART_GRAPHIC_URI}">'
        '<c:chart r:id="rId999"/>'
        "</a:graphicData>"
    )


def _picture_graphic_data(embed_id: str) -> str:
    return (
        f'<a:graphicData uri="{PICTURE_GRAPHIC_URI}">'
        "<pic:pic>"
        '<pic:nvPicPr><pic:cNvPr id="0" name="p"/><pic:cNvPicPr/></pic:nvPicPr>'
        f'<pic:blipFill><a:blip r:embed="{embed_id}"/></pic:blipFill>'
        "<pic:spPr/>"
        "</pic:pic>"
        "</a:graphicData>"
    )


def _drawing_xml(graphic_data: str, with_extent: bool = True) -> str:
    extent = '<wp:extent cx="914400" cy="457200"/>' if with_extent else ""
    return (
        "<w:drawing><wp:inline>"
        f"{extent}"
        f"<a:graphic>{graphic_data}</a:graphic>"
        "</wp:inline></w:drawing>"
    )


def _paragraph_xml(*runs: str) -> str:
    body = "".join(f"<w:r>{run}</w:r>" for run in runs)
    return f'<w:p {nsdecls("w", "wp", "a", "r", "c", "pic")}>{body}</w:p>'


def _append_paragraph(document, *runs: str) -> None:
    document.element.body.get_or_add_sectPr().addprevious(
        parse_xml(_paragraph_xml(*runs))
    )


def _add_emf_part(document) -> str:
    """Relate a fake EMF image part and return its relationship id."""
    part = Part(
        PackURI("/word/media/image99.emf"),
        "image/x-emf",
        b"\x01\x00\x00\x00EMF-bytes",
        document.part.package,
    )
    return document.part.relate_to(part, RT.IMAGE)


def _png_file(tmp_path, name: str = "img.png", color: str = "magenta") -> str:
    path = str(tmp_path / name)
    Image.new("RGB", (20, 10), color).save(path, format="PNG")
    return path


def _graphic_uri(drawing) -> str:
    return drawing.find(f".//{qn('a:graphicData')}").get("uri")


# --- selection -------------------------------------------------------------


def test_selects_lone_chart_paragraph():
    document = Document()
    document.add_paragraph("Before the chart.")
    _append_paragraph(document, _drawing_xml(_chart_graphic_data()))
    document.add_paragraph("After the chart.")

    renderable = _renderable_paragraphs(document)

    assert len(renderable) == 1
    _, drawing = renderable[0]
    assert _graphic_uri(drawing) == CHART_GRAPHIC_URI


def test_chart_sharing_a_paragraph_with_text_is_not_renderable():
    document = Document()
    _append_paragraph(
        document, "<w:t>Figure 1 caption</w:t>", _drawing_xml(_chart_graphic_data())
    )

    assert _renderable_paragraphs(document) == []


def test_two_drawings_in_one_paragraph_are_not_renderable():
    document = Document()
    _append_paragraph(
        document,
        _drawing_xml(_chart_graphic_data()),
        _drawing_xml(_chart_graphic_data()),
    )

    assert _renderable_paragraphs(document) == []


def test_chart_inside_a_table_is_not_renderable():
    document = Document()
    table_xml = (
        f'<w:tbl {nsdecls("w")}><w:tr><w:tc>'
        + _paragraph_xml(_drawing_xml(_chart_graphic_data()))
        + "</w:tc></w:tr></w:tbl>"
    )
    document.element.body.get_or_add_sectPr().addprevious(parse_xml(table_xml))

    assert _renderable_paragraphs(document) == []


def test_regular_picture_is_not_renderable(tmp_path):
    document = Document()
    with open(_png_file(tmp_path), "rb") as f:
        document.add_picture(io.BytesIO(f.read()))

    assert _renderable_paragraphs(document) == []


def test_metafile_picture_is_renderable():
    document = Document()
    embed_id = _add_emf_part(document)
    _append_paragraph(document, _drawing_xml(_picture_graphic_data(embed_id)))

    renderable = _renderable_paragraphs(document)

    assert len(renderable) == 1
    _, drawing = renderable[0]
    assert _graphic_uri(drawing) == PICTURE_GRAPHIC_URI


# --- render-source construction --------------------------------------------


def test_render_source_keeps_only_renderable_paragraphs(tmp_path):
    document = Document()
    document.add_paragraph("PARA-0")
    _append_paragraph(document, _drawing_xml(_chart_graphic_data()))
    document.add_paragraph("PARA-1")
    _append_paragraph(document, _drawing_xml(_chart_graphic_data()))
    src, dst = str(tmp_path / "src.docx"), str(tmp_path / "dst.docx")
    document.save(src)

    kept = _write_render_source_docx(src, dst)

    assert kept == 2
    rendered = Document(dst)
    body_paragraphs = [
        el for el in rendered.element.body if el.tag == qn("w:p")
    ]
    assert len(body_paragraphs) == 2
    assert "PARA" not in "\n".join(p.text for p in rendered.paragraphs)
    # Page breaks put each drawing on its own PDF page; the first needs none.
    breaks = [
        p.find(qn("w:pPr")) is not None
        and p.find(qn("w:pPr")).find(qn("w:pageBreakBefore")) is not None
        for p in body_paragraphs
    ]
    assert breaks == [False, True]
    # Page geometry survives.
    assert rendered.element.body.find(qn("w:sectPr")) is not None


def test_render_source_strips_headers_and_footers(tmp_path):
    document = Document()
    document.sections[0].header.paragraphs[0].text = "Running header"
    document.sections[0].footer.paragraphs[0].text = "Page footer"
    _append_paragraph(document, _drawing_xml(_chart_graphic_data()))
    src, dst = str(tmp_path / "src.docx"), str(tmp_path / "dst.docx")
    document.save(src)

    assert _write_render_source_docx(src, dst) == 1

    rendered = Document(dst)
    assert rendered.element.findall(f".//{qn('w:headerReference')}") == []
    assert rendered.element.findall(f".//{qn('w:footerReference')}") == []


# --- splicing rendered images back ------------------------------------------


def test_replace_drawing_swaps_chart_for_picture(tmp_path):
    document = Document()
    _append_paragraph(document, _drawing_xml(_chart_graphic_data()))
    _, drawing = _renderable_paragraphs(document)[0]

    assert _replace_drawing_with_image(document, drawing, _png_file(tmp_path))

    graphic_data = drawing.find(f".//{qn('a:graphicData')}")
    assert graphic_data.get("uri") == PICTURE_GRAPHIC_URI
    blip = graphic_data.find(f".//{qn('a:blip')}")
    embed_id = blip.get(qn("r:embed"))
    assert document.part.related_parts[embed_id].content_type == "image/png"
    # The drawing keeps the size the document declared.
    ext = graphic_data.find(f".//{qn('a:ext')}")
    assert (ext.get("cx"), ext.get("cy")) == ("914400", "457200")


def test_replace_drawing_without_extent_fails(tmp_path):
    document = Document()
    _append_paragraph(document, _drawing_xml(_chart_graphic_data(), with_extent=False))
    drawing = next(document.element.body.iter(qn("w:drawing")))

    assert not _replace_drawing_with_image(document, drawing, _png_file(tmp_path))


# --- orchestration -----------------------------------------------------------


@pytest.mark.asyncio
async def test_rasterize_returns_none_without_render_targets(tmp_path):
    document = Document()
    document.add_paragraph("Only text.")
    src = str(tmp_path / "doc.docx")
    document.save(src)

    with patch(f"{MODULE}._render_drawings_to_images") as render_mock:
        assert await rasterize_docx_drawings(src) is None

    render_mock.assert_not_called()


@pytest.mark.asyncio
async def test_rasterize_replaces_drawings_with_rendered_images(tmp_path):
    document = Document()
    document.add_paragraph("PARA-0")
    _append_paragraph(document, _drawing_xml(_chart_graphic_data()))
    _append_paragraph(document, _drawing_xml(_chart_graphic_data()))
    src = str(tmp_path / "doc.docx")
    document.save(src)

    render_dir = tempfile.mkdtemp()  # removed by the function under test
    renders = [
        _png_file(tmp_path, "r0.png", "magenta"),
        _png_file(tmp_path, "r1.png", "navy"),
    ]
    with patch(
        f"{MODULE}._render_drawings_to_images",
        new=AsyncMock(return_value=(render_dir, renders)),
    ):
        out = await rasterize_docx_drawings(src)

    assert out is not None
    try:
        rendered = Document(out)
        uris = [
            gd.get("uri")
            for gd in rendered.element.body.iter(qn("a:graphicData"))
        ]
        assert uris == [PICTURE_GRAPHIC_URI, PICTURE_GRAPHIC_URI]
        assert "PARA-0" in "\n".join(p.text for p in rendered.paragraphs)
        assert not os.path.exists(render_dir)
    finally:
        os.remove(out)


@pytest.mark.asyncio
async def test_rasterize_swallows_failures(tmp_path):
    broken = str(tmp_path / "not-a-docx.docx")
    with open(broken, "wb") as f:
        f.write(b"garbage")

    assert await rasterize_docx_drawings(broken) is None

# --- render-target detection edge cases ---------------------------------------


def test_drawing_without_graphic_data_is_not_renderable():
    document = Document()
    _append_paragraph(
        document,
        '<w:drawing><wp:inline><wp:extent cx="9525" cy="9525"/>'
        "<a:graphic/></wp:inline></w:drawing>",
    )

    assert _renderable_paragraphs(document) == []


def test_unknown_graphic_uri_is_not_renderable():
    document = Document()
    _append_paragraph(
        document,
        _drawing_xml('<a:graphicData uri="http://example.com/diagram"/>'),
    )

    assert _renderable_paragraphs(document) == []


def test_picture_without_blip_embed_is_not_renderable():
    document = Document()
    graphic_data = (
        f'<a:graphicData uri="{PICTURE_GRAPHIC_URI}">'
        "<pic:pic><pic:blipFill><a:blip/></pic:blipFill></pic:pic>"
        "</a:graphicData>"
    )
    _append_paragraph(document, _drawing_xml(graphic_data))

    assert _renderable_paragraphs(document) == []


def test_picture_with_dangling_relationship_is_not_renderable():
    document = Document()
    _append_paragraph(document, _drawing_xml(_picture_graphic_data("rId404")))

    assert _renderable_paragraphs(document) == []


# --- rendering-failure guards --------------------------------------------------


def _chart_docx(tmp_path) -> str:
    document = Document()
    _append_paragraph(document, _drawing_xml(_chart_graphic_data()))
    path = str(tmp_path / "chart.docx")
    document.save(path)
    return path


@pytest.mark.asyncio
async def test_rasterize_returns_none_when_rendering_fails(tmp_path):
    with patch(
        f"{MODULE}._render_drawings_to_images", new=AsyncMock(return_value=None)
    ):
        assert await rasterize_docx_drawings(_chart_docx(tmp_path)) is None


@pytest.mark.asyncio
async def test_rasterize_returns_none_when_nothing_could_be_replaced(tmp_path):
    """A renderable drawing whose splice fails must not produce a document
    identical to the original under a new name."""
    render_dir = tempfile.mkdtemp()  # removed by the function under test
    with (
        patch(
            f"{MODULE}._render_drawings_to_images",
            new=AsyncMock(return_value=(render_dir, [_png_file(tmp_path)])),
        ),
        patch(f"{MODULE}._replace_drawing_with_image", return_value=False),
    ):
        assert await rasterize_docx_drawings(_chart_docx(tmp_path)) is None
    assert not os.path.exists(render_dir)


@pytest.mark.asyncio
async def test_render_returns_none_without_libreoffice(tmp_path):
    with patch(f"{MODULE}.shutil.which", return_value=None):
        assert await rasterize_docx_drawings(_chart_docx(tmp_path)) is None


@pytest.mark.asyncio
async def test_render_bails_on_render_source_count_mismatch(tmp_path):
    with (
        patch(f"{MODULE}.shutil.which", return_value="/usr/bin/soffice"),
        patch(f"{MODULE}._write_render_source_docx", return_value=0),
    ):
        assert await rasterize_docx_drawings(_chart_docx(tmp_path)) is None


def _fake_process(returncode: int = 0, communicate=None):
    process = MagicMock()
    process.returncode = returncode
    process.communicate = communicate or AsyncMock(return_value=(b"", b"boom"))
    process.kill = MagicMock()
    process.wait = AsyncMock()
    return process


@pytest.mark.asyncio
async def test_render_bails_when_libreoffice_fails(tmp_path):
    with (
        patch(f"{MODULE}.shutil.which", return_value="/usr/bin/soffice"),
        patch(f"{MODULE}._write_render_source_docx", return_value=1),
        patch(
            f"{MODULE}.asyncio.create_subprocess_exec",
            new=AsyncMock(return_value=_fake_process(returncode=1)),
        ),
    ):
        assert await rasterize_docx_drawings(_chart_docx(tmp_path)) is None


@pytest.mark.asyncio
async def test_render_bails_on_page_count_mismatch(tmp_path):
    def write_source_and_pdf(src_path, dst_path):
        # Stand in for LibreOffice: the "PDF" appears next to the source.
        with open(os.path.join(os.path.dirname(dst_path), "render-source.pdf"), "wb") as f:
            f.write(b"%PDF-fake")
        return 1

    with (
        patch(f"{MODULE}.shutil.which", return_value="/usr/bin/soffice"),
        patch(
            f"{MODULE}._write_render_source_docx", side_effect=write_source_and_pdf
        ),
        patch(
            f"{MODULE}.asyncio.create_subprocess_exec",
            new=AsyncMock(return_value=_fake_process(returncode=0)),
        ),
        patch(f"{MODULE}._render_pdf_pages", return_value=[]),
    ):
        assert await rasterize_docx_drawings(_chart_docx(tmp_path)) is None


@pytest.mark.asyncio
async def test_render_bails_on_libreoffice_timeout(tmp_path):
    process = _fake_process(communicate=AsyncMock(side_effect=asyncio.TimeoutError))
    with (
        patch(f"{MODULE}.shutil.which", return_value="/usr/bin/soffice"),
        patch(f"{MODULE}._write_render_source_docx", return_value=1),
        patch(
            f"{MODULE}.asyncio.create_subprocess_exec",
            new=AsyncMock(return_value=process),
        ),
    ):
        assert await rasterize_docx_drawings(_chart_docx(tmp_path)) is None

    process.kill.assert_called_once()


@pytest.mark.asyncio
async def test_replace_drawing_survives_image_part_failure(tmp_path):
    document = Document()
    _append_paragraph(document, _drawing_xml(_chart_graphic_data()))
    _, drawing = _renderable_paragraphs(document)[0]

    with patch.object(
        type(document.part), "get_or_add_image", side_effect=RuntimeError("disk full")
    ):
        assert not _replace_drawing_with_image(document, drawing, _png_file(tmp_path))

    # The drawing is left untouched for the caller's replaced-count guard.
    assert _graphic_uri(drawing) == CHART_GRAPHIC_URI


def test_chart_without_extent_is_not_renderable():
    """No declared size means nothing to render at or crop to."""
    document = Document()
    _append_paragraph(document, _drawing_xml(_chart_graphic_data(), with_extent=False))

    assert _renderable_paragraphs(document) == []


def test_render_source_pins_drawings_to_page_origin(tmp_path):
    """Zero margins/spacing/indent make the geometric crop valid: the drawing
    occupies exactly its declared extent from the page's top-left corner."""
    document = Document()
    _append_paragraph(document, _drawing_xml(_chart_graphic_data()))
    src, dst = str(tmp_path / "src.docx"), str(tmp_path / "dst.docx")
    document.save(src)

    assert _write_render_source_docx(src, dst) == 1

    rendered = Document(dst)
    margins = rendered.element.body.find(qn("w:sectPr")).find(qn("w:pgMar"))
    assert {margins.get(qn(f"w:{a}")) for a in ("top", "right", "bottom", "left")} == {"0"}
    paragraph = next(el for el in rendered.element.body if el.tag == qn("w:p"))
    pPr = paragraph.find(qn("w:pPr"))
    spacing = pPr.find(qn("w:spacing"))
    assert (spacing.get(qn("w:before")), spacing.get(qn("w:after"))) == ("0", "0")
    assert pPr.find(qn("w:ind")).get(qn("w:left")) == "0"
    # Centered figures would otherwise sit mid-page, outside the crop.
    assert pPr.find(qn("w:jc")).get(qn("w:val")) == "left"


def test_render_pdf_pages_crops_to_declared_extent(tmp_path):
    """Pages are cropped geometrically to each drawing's extent — the crop
    must not depend on what is visible (a blank drawing crops the same)."""
    # A 2-page PDF of 200x100pt pages (PIL writes PDFs at 72 dpi).
    pages = [
        Image.new("RGB", (200, 100), "white"),
        Image.new("RGB", (200, 100), "magenta"),
    ]
    pdf_path = str(tmp_path / "pages.pdf")
    pages[0].save(pdf_path, save_all=True, append_images=pages[1:])

    extents = [
        (150 * 12700, 50 * 12700),  # 150x50pt
        (400 * 12700, 400 * 12700),  # larger than the page: clamped
    ]
    paths = _render_pdf_pages(pdf_path, str(tmp_path), extents)

    scale = 4
    first, second = (Image.open(p) for p in paths)
    assert first.size == (150 * scale, 50 * scale)
    assert second.size == (200 * scale, 100 * scale)  # clamped to the page


def test_render_pdf_pages_signals_page_count_mismatch(tmp_path):
    page = Image.new("RGB", (100, 100), "white")
    pdf_path = str(tmp_path / "one-page.pdf")
    page.save(pdf_path)

    assert _render_pdf_pages(pdf_path, str(tmp_path), [(9525, 9525)] * 2) == []


def test_render_source_creates_page_margins_when_absent(tmp_path):
    document = Document()
    _append_paragraph(document, _drawing_xml(_chart_graphic_data()))
    sectPr = document.element.body.get_or_add_sectPr()
    margins = sectPr.find(qn("w:pgMar"))
    if margins is not None:
        sectPr.remove(margins)
    src, dst = str(tmp_path / "src.docx"), str(tmp_path / "dst.docx")
    document.save(src)

    assert _write_render_source_docx(src, dst) == 1

    rendered_margins = Document(dst).element.body.find(qn("w:sectPr")).find(qn("w:pgMar"))
    assert rendered_margins.get(qn("w:top")) == "0"


def _anchored_drawing_xml(graphic_data: str) -> str:
    return (
        "<w:drawing><wp:anchor>"
        '<wp:positionH relativeFrom="page"><wp:posOffset>1000000</wp:posOffset></wp:positionH>'
        '<wp:positionV relativeFrom="page"><wp:posOffset>2000000</wp:posOffset></wp:positionV>'
        '<wp:extent cx="914400" cy="457200"/>'
        f"<a:graphic>{graphic_data}</a:graphic>"
        "</wp:anchor></w:drawing>"
    )


def test_render_source_converts_anchored_drawings_to_inline(tmp_path):
    """A floating drawing keeps its position offsets, which would move it
    away from the page origin the geometric crop assumes — the render source
    makes it inline so the pinned paragraph places it at (0, 0)."""
    document = Document()
    _append_paragraph(document, _anchored_drawing_xml(_chart_graphic_data()))
    assert len(_renderable_paragraphs(document)) == 1  # anchors are selected
    src, dst = str(tmp_path / "src.docx"), str(tmp_path / "dst.docx")
    document.save(src)

    assert _write_render_source_docx(src, dst) == 1

    rendered = Document(dst)
    assert rendered.element.findall(f".//{qn('wp:anchor')}") == []
    inline = rendered.element.body.find(f".//{qn('wp:inline')}")
    extent = inline.find(qn("wp:extent"))
    assert (extent.get("cx"), extent.get("cy")) == ("914400", "457200")
    assert inline.find(f".//{qn('a:graphicData')}").get("uri") == CHART_GRAPHIC_URI


@pytest.mark.asyncio
async def test_rasterize_is_all_or_nothing_on_partial_splice_failure(tmp_path):
    """A transient splice failure must not ship a half-replaced document —
    the next run could replace a different subset and break line parity."""
    document = Document()
    _append_paragraph(document, _drawing_xml(_chart_graphic_data()))
    _append_paragraph(document, _drawing_xml(_chart_graphic_data()))
    src = str(tmp_path / "two-charts.docx")
    document.save(src)

    render_dir = tempfile.mkdtemp()  # removed by the function under test
    renders = [_png_file(tmp_path, "r0.png"), _png_file(tmp_path, "r1.png")]
    with (
        patch(
            f"{MODULE}._render_drawings_to_images",
            new=AsyncMock(return_value=(render_dir, renders)),
        ),
        patch(
            f"{MODULE}._replace_drawing_with_image", side_effect=[True, False]
        ),
    ):
        assert await rasterize_docx_drawings(src) is None
    assert not os.path.exists(render_dir)


@pytest.mark.asyncio
async def test_render_cleans_tmp_dir_on_unexpected_failure(tmp_path):
    """Exceptions beyond the modeled ones (subprocess spawn errors, PDFium
    crashes) must not strand drawing-raster-* directories."""
    render_dir = str(tmp_path / "raster-tmp")
    os.makedirs(render_dir)

    with (
        patch(f"{MODULE}.shutil.which", return_value="/usr/bin/soffice"),
        patch(f"{MODULE}.tempfile.mkdtemp", return_value=render_dir),
        patch(f"{MODULE}._write_render_source_docx", return_value=1),
        patch(
            f"{MODULE}.asyncio.create_subprocess_exec",
            new=AsyncMock(side_effect=OSError("spawn failed")),
        ),
    ):
        assert await rasterize_docx_drawings(_chart_docx(tmp_path)) is None

    assert not os.path.exists(render_dir)
