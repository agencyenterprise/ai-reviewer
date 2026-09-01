"""Unit tests for reading image placements (size and crop) out of a DOCX."""

import io

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches
from PIL import Image
from xxhash import xxh128

from lib.services.docx.image_display_sizes import (
    ImagePlacement,
    SourceRect,
    read_docx_image_display_sizes,
)


def _png_bytes(color: str) -> bytes:
    buf = io.BytesIO()
    Image.new("RGB", (300, 100), color).save(buf, format="PNG")
    return buf.getvalue()


def test_reads_display_size_in_css_pixels(tmp_path):
    png = _png_bytes("magenta")
    doc = Document()
    doc.add_picture(io.BytesIO(png), width=Inches(2), height=Inches(1))
    path = str(tmp_path / "doc.docx")
    doc.save(path)

    sizes = read_docx_image_display_sizes(path)

    # 96 px per inch, and no crop declared.
    assert sizes.take(xxh128(png).hexdigest()) == ImagePlacement(
        width_px=192, height_px=96
    )


def test_same_image_twice_yields_sizes_in_document_order(tmp_path):
    png = _png_bytes("navy")
    doc = Document()
    doc.add_picture(io.BytesIO(png), width=Inches(2), height=Inches(1))
    doc.add_picture(io.BytesIO(png), width=Inches(1), height=Inches(1))
    path = str(tmp_path / "doc.docx")
    doc.save(path)

    sizes = read_docx_image_display_sizes(path)
    content_hash = xxh128(png).hexdigest()

    assert sizes.take(content_hash) == ImagePlacement(width_px=192, height_px=96)
    assert sizes.take(content_hash) == ImagePlacement(width_px=96, height_px=96)
    assert sizes.take(content_hash) is None


def test_unreadable_docx_returns_empty_sizes(tmp_path):
    path = str(tmp_path / "not-a-docx.docx")
    with open(path, "wb") as f:
        f.write(b"garbage")

    sizes = read_docx_image_display_sizes(path)

    assert sizes.take("anything") is None


def test_reads_size_of_anchored_floating_image(tmp_path):
    """Floating images (wp:anchor) are not in `document.inline_shapes`; sizes
    must come from walking the drawing XML — a cover-page logo regression."""
    png = _png_bytes("magenta")
    path = str(tmp_path / "img.png")
    with open(path, "wb") as f:
        f.write(png)

    doc = Document()
    embed_id, _ = doc.part.get_or_add_image(path)
    drawing = (
        f'<w:p {nsdecls("w", "wp", "a", "r", "pic")}><w:r><w:drawing>'
        "<wp:anchor>"
        '<wp:extent cx="580644" cy="580644"/>'
        '<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic><pic:nvPicPr><pic:cNvPr id="0" name="logo"/><pic:cNvPicPr/></pic:nvPicPr>'
        f'<pic:blipFill><a:blip r:embed="{embed_id}"/></pic:blipFill><pic:spPr/></pic:pic>'
        "</a:graphicData></a:graphic>"
        "</wp:anchor></w:drawing></w:r></w:p>"
    )
    doc.element.body.get_or_add_sectPr().addprevious(parse_xml(drawing))
    docx_path = str(tmp_path / "doc.docx")
    doc.save(docx_path)

    sizes = read_docx_image_display_sizes(docx_path)

    assert sizes.take(xxh128(png).hexdigest()) == ImagePlacement(
        width_px=61, height_px=61
    )


def test_malformed_drawings_are_skipped(tmp_path):
    """Drawings without an extent, without a blip, or with a dangling
    relationship id must be skipped rather than fail the read."""
    doc = Document()
    paragraphs = [
        # No extent.
        "<w:drawing><wp:inline><a:graphic><a:graphicData "
        'uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic><pic:blipFill><a:blip r:embed="rId9"/></pic:blipFill></pic:pic>'
        "</a:graphicData></a:graphic></wp:inline></w:drawing>",
        # No blip at all.
        '<w:drawing><wp:inline><wp:extent cx="9525" cy="9525"/>'
        "<a:graphic><a:graphicData "
        'uri="http://schemas.openxmlformats.org/drawingml/2006/picture"/>'
        "</a:graphic></wp:inline></w:drawing>",
        # Blip without an embed id.
        '<w:drawing><wp:inline><wp:extent cx="9525" cy="9525"/>'
        "<a:graphic><a:graphicData "
        'uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        "<pic:pic><pic:blipFill><a:blip/></pic:blipFill></pic:pic>"
        "</a:graphicData></a:graphic></wp:inline></w:drawing>",
        # Dangling relationship id.
        '<w:drawing><wp:inline><wp:extent cx="9525" cy="9525"/>'
        "<a:graphic><a:graphicData "
        'uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic><pic:blipFill><a:blip r:embed="rId404"/></pic:blipFill></pic:pic>'
        "</a:graphicData></a:graphic></wp:inline></w:drawing>",
    ]
    for inner in paragraphs:
        xml = f'<w:p {nsdecls("w", "wp", "a", "r", "pic")}><w:r>{inner}</w:r></w:p>'
        doc.element.body.get_or_add_sectPr().addprevious(parse_xml(xml))
    path = str(tmp_path / "doc.docx")
    doc.save(path)

    sizes = read_docx_image_display_sizes(path)

    assert sizes.take("any-hash") is None


def _docx_with_src_rect(tmp_path, png: bytes, src_rect: str) -> str:
    """A one-picture DOCX whose blipFill carries ``src_rect`` verbatim."""
    image_path = str(tmp_path / "img.png")
    with open(image_path, "wb") as f:
        f.write(png)

    doc = Document()
    embed_id, _ = doc.part.get_or_add_image(image_path)
    xml = (
        f'<w:p {nsdecls("w", "wp", "a", "r", "pic")}><w:r><w:drawing><wp:inline>'
        '<wp:extent cx="1905000" cy="952500"/>'
        '<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic><pic:nvPicPr><pic:cNvPr id="0" name="fig"/><pic:cNvPicPr/></pic:nvPicPr>'
        f'<pic:blipFill><a:blip r:embed="{embed_id}"/>{src_rect}</pic:blipFill>'
        "<pic:spPr/></pic:pic>"
        "</a:graphicData></a:graphic>"
        "</wp:inline></w:drawing></w:r></w:p>"
    )
    doc.element.body.get_or_add_sectPr().addprevious(parse_xml(xml))
    docx_path = str(tmp_path / "cropped.docx")
    doc.save(docx_path)
    return docx_path


def test_reads_declared_crop(tmp_path):
    """The extent describes the *cropped* region, so the crop has to come
    along with it — reading one without the other squashes the figure."""
    png = _png_bytes("teal")
    path = _docx_with_src_rect(tmp_path, png, '<a:srcRect t="17222" b="3597"/>')

    placement = read_docx_image_display_sizes(path).take(xxh128(png).hexdigest())

    assert placement == ImagePlacement(
        width_px=200,
        height_px=100,
        crop=SourceRect(top=0.17222, bottom=0.03597),
    )


def test_percentage_string_crop_is_understood(tmp_path):
    """The strict schema writes ST_Percentage as a literal percentage."""
    png = _png_bytes("olive")
    path = _docx_with_src_rect(tmp_path, png, '<a:srcRect l="12.5%" r="25%"/>')

    placement = read_docx_image_display_sizes(path).take(xxh128(png).hexdigest())

    assert placement is not None
    assert placement.crop == SourceRect(left=0.125, right=0.25)


def test_empty_src_rect_reads_as_uncropped(tmp_path):
    """Word writes an empty srcRect for uncropped pictures."""
    png = _png_bytes("purple")
    path = _docx_with_src_rect(tmp_path, png, "<a:srcRect/>")

    placement = read_docx_image_display_sizes(path).take(xxh128(png).hexdigest())

    assert placement is not None
    assert placement.crop is None


def test_outset_edges_keep_their_sign(tmp_path):
    """Negative edges are outsets (Word pads that side). Clamping them to zero
    would silently reshape the region: `l="-10000" r="10000"` is a full-width
    band shifted left, not a 90%-wide crop, so the signs have to survive the
    read for `has_outset` to catch it."""
    png = _png_bytes("purple")
    path = _docx_with_src_rect(tmp_path, png, '<a:srcRect l="-10000" r="10000"/>')

    placement = read_docx_image_display_sizes(path).take(xxh128(png).hexdigest())

    assert placement is not None
    assert placement.crop == SourceRect(left=-0.1, right=0.1)
    assert not placement.crop.is_applicable


def test_unparseable_crop_edge_is_reported_as_unreadable(tmp_path):
    """Reading a malformed edge as zero would report a cropped picture as
    uncropped, and the declared size would then stretch the whole image."""
    png = _png_bytes("maroon")
    path = _docx_with_src_rect(tmp_path, png, '<a:srcRect t="lots"/>')

    placement = read_docx_image_display_sizes(path).take(xxh128(png).hexdigest())

    assert placement is not None
    assert placement.crop == SourceRect(unreadable=True)
    assert not placement.crop.is_applicable


def test_missing_and_empty_edges_are_simply_untrimmed(tmp_path):
    """Absent attributes are not parse failures — only present-but-malformed
    ones are, or every partially-cropped picture would be unreadable."""
    png = _png_bytes("silver")
    path = _docx_with_src_rect(tmp_path, png, '<a:srcRect t="25000" b=""/>')

    placement = read_docx_image_display_sizes(path).take(xxh128(png).hexdigest())

    assert placement is not None
    assert placement.crop == SourceRect(top=0.25)
    assert placement.crop.is_applicable


def test_crop_box_is_none_when_nothing_would_change():
    """A sub-pixel crop must not force callers to re-encode."""
    assert SourceRect(top=0.001, bottom=0.001).crop_box(100, 100) is None
    assert SourceRect().crop_box(100, 100) is None
    assert SourceRect(top=0.25, bottom=0.25).crop_box(100, 100) == (0, 25, 100, 75)


def test_is_applicable_separates_a_real_crop_from_the_three_that_are_not():
    """`crop_box` returning None only ever means "nothing to trim", so the
    regions that cannot be cropped at all have to be caught before it —
    otherwise they read as a no-op and keep their declared size."""
    assert SourceRect(top=0.25, bottom=0.1).is_applicable
    assert SourceRect().is_applicable
    assert not SourceRect(unreadable=True).is_applicable
    assert not SourceRect(left=-0.1, right=0.1).is_applicable  # outset
    assert not SourceRect(left=0.6, right=0.6).is_applicable  # insets overlap
    assert not SourceRect(top=0.5, bottom=0.5).is_applicable  # exactly nothing
